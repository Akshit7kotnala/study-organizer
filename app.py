import os
from datetime import datetime
from uuid import uuid4

from flask import (
    Flask,
    render_template,
    request,
    redirect,
    url_for,
    send_from_directory,
    flash,
    session,
    jsonify,
)
from flask_sqlalchemy import SQLAlchemy
from flask_login import LoginManager, UserMixin, login_user, logout_user, login_required, current_user
from authlib.integrations.flask_client import OAuth
from werkzeug.utils import secure_filename
from dotenv import load_dotenv
from PIL import Image
from pdf2image import convert_from_path
import boto3
from botocore.exceptions import ClientError
from azure.storage.blob import BlobServiceClient, generate_blob_sas, BlobSasPermissions
from azure.core.exceptions import AzureError
from datetime import timedelta
import openai
from openai import OpenAI
import google.generativeai as genai
import PyPDF2
import pytesseract
from docx import Document as DocxDocument
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import nltk
from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize
import re
import io

# Load environment variables from .env file
load_dotenv()

basedir = os.path.abspath(os.path.dirname(__file__))

app = Flask(__name__)
app.config['SECRET_KEY'] = os.environ.get('SECRET_KEY', 'dev-secret-change-me-in-production')

# Database configuration - use PostgreSQL in production, SQLite in development
database_url = os.environ.get('DATABASE_URL')
if database_url:
    # Fix for Render.com PostgreSQL URL format
    if database_url.startswith('postgres://'):
        database_url = database_url.replace('postgres://', 'postgresql://', 1)
    app.config['SQLALCHEMY_DATABASE_URI'] = database_url
else:
    # Development - use SQLite
    app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///' + os.path.join(basedir, 'documents.db')

app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
app.config['UPLOAD_FOLDER'] = os.path.join(basedir, 'uploads')
app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024  # 50MB max file size

# Google OAuth config - get these from Google Cloud Console
app.config['GOOGLE_CLIENT_ID'] = os.environ.get('GOOGLE_CLIENT_ID', '')
app.config['GOOGLE_CLIENT_SECRET'] = os.environ.get('GOOGLE_CLIENT_SECRET', '')

# AWS S3 Configuration
app.config['AWS_ACCESS_KEY_ID'] = os.environ.get('AWS_ACCESS_KEY_ID', '')
app.config['AWS_SECRET_ACCESS_KEY'] = os.environ.get('AWS_SECRET_ACCESS_KEY', '')
app.config['AWS_REGION'] = os.environ.get('AWS_REGION', 'us-east-1')
app.config['S3_BUCKET_NAME'] = os.environ.get('S3_BUCKET_NAME', '')

# Azure Blob Storage Configuration
app.config['AZURE_STORAGE_ACCOUNT_NAME'] = os.environ.get('AZURE_STORAGE_ACCOUNT_NAME', '')
app.config['AZURE_STORAGE_ACCOUNT_KEY'] = os.environ.get('AZURE_STORAGE_ACCOUNT_KEY', '')
app.config['AZURE_CONTAINER_NAME'] = os.environ.get('AZURE_CONTAINER_NAME', 'study-documents')

# Storage Configuration: 'local', 's3', or 'azure'
app.config['STORAGE_TYPE'] = os.environ.get('STORAGE_TYPE', 'local')

os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

# Initialize S3 client if using S3 storage
s3_client = None
if app.config['STORAGE_TYPE'] == 's3' and app.config['AWS_ACCESS_KEY_ID']:
    try:
        s3_client = boto3.client(
            's3',
            aws_access_key_id=app.config['AWS_ACCESS_KEY_ID'],
            aws_secret_access_key=app.config['AWS_SECRET_ACCESS_KEY'],
            region_name=app.config['AWS_REGION']
        )
        print(f"✓ S3 client initialized for bucket: {app.config['S3_BUCKET_NAME']}")
    except Exception as e:
        print(f"✗ Failed to initialize S3 client: {e}")
        print("  Falling back to local storage")
        app.config['STORAGE_TYPE'] = 'local'

# Initialize Azure Blob Service Client if using Azure storage
blob_service_client = None
if app.config['STORAGE_TYPE'] == 'azure' and app.config['AZURE_STORAGE_ACCOUNT_NAME']:
    try:
        connection_string = f"DefaultEndpointsProtocol=https;AccountName={app.config['AZURE_STORAGE_ACCOUNT_NAME']};AccountKey={app.config['AZURE_STORAGE_ACCOUNT_KEY']};EndpointSuffix=core.windows.net"
        blob_service_client = BlobServiceClient.from_connection_string(connection_string)
        
        # Create container if it doesn't exist
        container_client = blob_service_client.get_container_client(app.config['AZURE_CONTAINER_NAME'])
        try:
            container_client.create_container()
            print(f"✓ Created Azure container: {app.config['AZURE_CONTAINER_NAME']}")
        except:
            pass  # Container already exists
        
        print(f"✓ Azure Blob Storage initialized for account: {app.config['AZURE_STORAGE_ACCOUNT_NAME']}")
    except Exception as e:
        print(f"✗ Failed to initialize Azure Blob Storage: {e}")
        print("  Falling back to local storage")
        app.config['STORAGE_TYPE'] = 'local'

# ============================================================================
# AI Configuration
# ============================================================================

# Google Gemini Configuration (FREE - Primary AI Provider)
app.config['GEMINI_API_KEY'] = os.environ.get('GEMINI_API_KEY', '')
gemini_model = None
if app.config['GEMINI_API_KEY']:
    try:
        genai.configure(api_key=app.config['GEMINI_API_KEY'])
        gemini_model = genai.GenerativeModel('gemini-2.0-flash')
        print(f"✓ Google Gemini AI initialized")
    except Exception as e:
        print(f"⚠ Failed to initialize Gemini: {e}")
        gemini_model = None
else:
    print(f"⚠ Gemini API key not found - AI features will be disabled")

# OpenAI Configuration (Deprecated - keeping for backward compatibility)
app.config['OPENAI_API_KEY'] = os.environ.get('OPENAI_API_KEY', '')
if app.config['OPENAI_API_KEY'] and app.config['OPENAI_API_KEY'] != 'your-openai-api-key-here':
    openai_client = OpenAI(api_key=app.config['OPENAI_API_KEY'])
    print(f"✓ OpenAI client initialized (fallback)")
else:
    openai_client = None

# Tesseract OCR Configuration
app.config['TESSERACT_CMD'] = os.environ.get('TESSERACT_CMD', 'C:\\Program Files\\Tesseract-OCR\\tesseract.exe')
if os.path.exists(app.config['TESSERACT_CMD']):
    pytesseract.pytesseract.tesseract_cmd = app.config['TESSERACT_CMD']
    print(f"✓ Tesseract OCR configured")
else:
    print(f"⚠ Tesseract OCR not found at {app.config['TESSERACT_CMD']} - OCR features will be disabled")

# AI Features Configuration
app.config['AI_SUMMARY_MAX_LENGTH'] = int(os.environ.get('AI_SUMMARY_MAX_LENGTH', 500))
app.config['AI_TAGS_COUNT'] = int(os.environ.get('AI_TAGS_COUNT', 5))
app.config['OCR_LANGUAGE'] = os.environ.get('OCR_LANGUAGE', 'eng')
app.config['SEARCH_RESULTS_LIMIT'] = int(os.environ.get('SEARCH_RESULTS_LIMIT', 50))
app.config['RECOMMENDATIONS_COUNT'] = int(os.environ.get('RECOMMENDATIONS_COUNT', 5))

# Initialize NLTK stopwords
try:
    stop_words = set(stopwords.words('english'))
except:
    stop_words = set()

db = SQLAlchemy(app)
login_manager = LoginManager(app)
login_manager.login_view = 'login'
login_manager.login_message = 'Please log in to access this page.'

oauth = OAuth(app)
google = oauth.register(
    name='google',
    client_id=app.config['GOOGLE_CLIENT_ID'],
    client_secret=app.config['GOOGLE_CLIENT_SECRET'],
    server_metadata_url='https://accounts.google.com/.well-known/openid-configuration',
    client_kwargs={'scope': 'openid email profile'}
)


# Models
# Association table for many-to-many relationship between Document and Tag
document_tags = db.Table('document_tags',
    db.Column('document_id', db.Integer, db.ForeignKey('document.id'), primary_key=True),
    db.Column('tag_id', db.Integer, db.ForeignKey('tag.id'), primary_key=True)
)

# Association table for many-to-many relationship between Document and Collection
document_collections = db.Table('document_collections',
    db.Column('document_id', db.Integer, db.ForeignKey('document.id'), primary_key=True),
    db.Column('collection_id', db.Integer, db.ForeignKey('collection.id'), primary_key=True)
)

# Association table for many-to-many relationship between StudyGroup and User
group_members = db.Table('group_members',
    db.Column('group_id', db.Integer, db.ForeignKey('study_group.id'), primary_key=True),
    db.Column('user_id', db.Integer, db.ForeignKey('user.id'), primary_key=True),
    db.Column('role', db.String(16), default='member'),  # 'admin', 'member'
    db.Column('joined_at', db.DateTime, default=datetime.utcnow)
)


class User(UserMixin, db.Model):
    id = db.Column(db.Integer, primary_key=True)
    email = db.Column(db.String(255), unique=True, nullable=False)
    name = db.Column(db.String(255), nullable=True)
    google_id = db.Column(db.String(255), unique=True, nullable=False)
    profile_pic = db.Column(db.String(512), nullable=True)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    
    documents = db.relationship('Document', backref='owner', lazy=True)
    shared_with_me = db.relationship('SharePermission', foreign_keys='SharePermission.shared_with_id', backref='recipient', lazy=True)
    my_shares = db.relationship('SharePermission', foreign_keys='SharePermission.shared_by_id', backref='sharer', lazy=True)
    comments = db.relationship('Comment', backref='author', lazy=True)
    notifications = db.relationship('Notification', backref='user', lazy=True, order_by='Notification.created_at.desc()')



class Tag(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(64), unique=True, nullable=False, index=True)
    slug = db.Column(db.String(64), unique=True, nullable=False, index=True)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    
    documents = db.relationship('Document', secondary=document_tags, back_populates='tag_objects')
    
    def __repr__(self):
        return f'<Tag {self.name}>'
    
    @staticmethod
    def slugify(text):
        """Convert tag name to URL-friendly slug."""
        return text.lower().strip().replace(' ', '-').replace('_', '-')


class Collection(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(128), nullable=False)
    description = db.Column(db.Text, nullable=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    color = db.Column(db.String(7), default='#667eea')  # Hex color for visual distinction
    icon = db.Column(db.String(32), default='bi-folder')  # Bootstrap icon class
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    updated_at = db.Column(db.DateTime, default=datetime.utcnow, onupdate=datetime.utcnow)
    
    # Many-to-many relationship with Document
    documents = db.relationship('Document', secondary=document_collections, backref='collections', lazy='dynamic')
    
    def __repr__(self):
        return f'<Collection {self.name}>'
    
    def document_count(self):
        """Return the number of documents in this collection."""
        return self.documents.count()


class Document(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    original_filename = db.Column(db.String(512), nullable=False)
    stored_filename = db.Column(db.String(512), nullable=False, unique=True)
    year = db.Column(db.Integer, nullable=False)
    subject = db.Column(db.String(128), nullable=False)
    tags = db.Column(db.String(256), nullable=True)  # Keep for backward compatibility
    mimetype = db.Column(db.String(128), nullable=True)
    size = db.Column(db.Integer, nullable=True)
    thumbnail_filename = db.Column(db.String(512), nullable=True)  # Thumbnail image
    storage_type = db.Column(db.String(16), default='local', nullable=False)  # 'local' or 's3'
    upload_date = db.Column(db.DateTime, default=datetime.utcnow)
    
    # AI-powered features
    summary = db.Column(db.Text, nullable=True)  # AI-generated summary
    extracted_text = db.Column(db.Text, nullable=True)  # Extracted text from PDF/images (OCR)
    ai_tags = db.Column(db.String(512), nullable=True)  # AI-suggested tags (comma-separated)
    content_vector = db.Column(db.Text, nullable=True)  # TF-IDF vector for recommendations (JSON)
    last_analyzed = db.Column(db.DateTime, nullable=True)  # Last AI analysis timestamp
    
    # Many-to-many relationship with Tag
    tag_objects = db.relationship('Tag', secondary=document_tags, back_populates='documents')

    def tag_list(self):
        """Return list of tag names (from tag_objects relationship)."""
        if self.tag_objects:
            return [tag.name for tag in self.tag_objects]
        # Fallback to old tags field if no tag_objects
        if not self.tags:
            return []
        return [t.strip() for t in self.tags.split(',') if t.strip()]
    
    def get_tags(self):
        """Return list of tag objects with name and slug."""
        if self.tag_objects:
            return [{'name': tag.name, 'slug': tag.slug} for tag in self.tag_objects]
        # Fallback for old comma-separated tags
        if not self.tags:
            return []
        return [{'name': t.strip(), 'slug': Tag.slugify(t.strip())} for t in self.tags.split(',') if t.strip()]
    
    def set_tags_from_string(self, tags_string):
        """Parse comma-separated tags and create/associate Tag objects."""
        if not tags_string or not tags_string.strip():
            self.tag_objects = []
            return
        
        tag_names = [t.strip().lower() for t in tags_string.split(',') if t.strip()]
        self.tag_objects = []
        
        for tag_name in tag_names:
            # Check if tag exists
            tag = Tag.query.filter_by(slug=Tag.slugify(tag_name)).first()
            if not tag:
                # Create new tag
                tag = Tag(name=tag_name, slug=Tag.slugify(tag_name))
                db.session.add(tag)
            self.tag_objects.append(tag)
    
    def format_size(self):
        """Return human-readable file size."""
        if not self.size:
            return 'Unknown'
        size = self.size
        for unit in ['B', 'KB', 'MB', 'GB']:
            if self.size < 1024.0:
                return f"{self.size:.1f} {unit}"
            self.size /= 1024.0
        return f"{self.size:.1f} TB"
    
    def file_icon(self):
        """Return Bootstrap icon class based on file type."""
        ext = os.path.splitext(self.original_filename)[1].lower()
        icon_map = {
            '.pdf': 'bi-file-earmark-pdf-fill text-danger',
            '.doc': 'bi-file-earmark-word-fill text-primary',
            '.docx': 'bi-file-earmark-word-fill text-primary',
            '.xls': 'bi-file-earmark-excel-fill text-success',
            '.xlsx': 'bi-file-earmark-excel-fill text-success',
            '.ppt': 'bi-file-earmark-ppt-fill text-warning',
            '.pptx': 'bi-file-earmark-ppt-fill text-warning',
            '.jpg': 'bi-file-earmark-image-fill text-info',
            '.jpeg': 'bi-file-earmark-image-fill text-info',
            '.png': 'bi-file-earmark-image-fill text-info',
            '.gif': 'bi-file-earmark-image-fill text-info',
            '.zip': 'bi-file-earmark-zip-fill text-secondary',
            '.txt': 'bi-file-earmark-text-fill',
        }
        return icon_map.get(ext, 'bi-file-earmark-fill')


class SharePermission(db.Model):
    """Model for sharing documents and collections with other users."""
    id = db.Column(db.Integer, primary_key=True)
    shared_by_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    shared_with_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    
    # What's being shared
    document_id = db.Column(db.Integer, db.ForeignKey('document.id'), nullable=True)
    collection_id = db.Column(db.Integer, db.ForeignKey('collection.id'), nullable=True)
    
    # Permission level: 'viewer', 'editor', 'admin'
    permission = db.Column(db.String(16), default='viewer', nullable=False)
    
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    
    # Relationships
    document = db.relationship('Document', backref='shares')
    collection = db.relationship('Collection', backref='shares')
    
    def __repr__(self):
        return f'<SharePermission {self.permission} for {self.document_id or self.collection_id}>'


class Comment(db.Model):
    """Model for comments/annotations on documents."""
    id = db.Column(db.Integer, primary_key=True)
    document_id = db.Column(db.Integer, db.ForeignKey('document.id'), nullable=False)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    content = db.Column(db.Text, nullable=False)
    page_number = db.Column(db.Integer, nullable=True)  # For PDFs, page annotation
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    updated_at = db.Column(db.DateTime, default=datetime.utcnow, onupdate=datetime.utcnow)
    
    # Relationships
    document = db.relationship('Document', backref='comments')
    
    def __repr__(self):
        return f'<Comment on Document {self.document_id} by User {self.user_id}>'


class StudyGroup(db.Model):
    """Model for study groups."""
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(128), nullable=False)
    description = db.Column(db.Text, nullable=True)
    created_by_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    icon = db.Column(db.String(32), default='bi-people-fill')
    color = db.Column(db.String(7), default='#667eea')
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    
    # Relationships
    creator = db.relationship('User', backref='created_groups')
    members = db.relationship('User', secondary=group_members, backref='study_groups')
    
    def __repr__(self):
        return f'<StudyGroup {self.name}>'
    
    def member_count(self):
        """Return the number of members in this group."""
        return len(self.members)
    
    def is_admin(self, user):
        """Check if user is admin of this group."""
        if user.id == self.created_by_id:
            return True
        # Could extend to check group_members role
        return False


class Notification(db.Model):
    """Model for user notifications."""
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    title = db.Column(db.String(256), nullable=False)
    message = db.Column(db.Text, nullable=False)
    type = db.Column(db.String(32), default='info')  # 'info', 'share', 'comment', 'group'
    link = db.Column(db.String(512), nullable=True)  # Link to relevant page
    read = db.Column(db.Boolean, default=False)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    
    def __repr__(self):
        return f'<Notification for User {self.user_id}>'


class ChatSession(db.Model):
    """Model for AI chat sessions with documents."""
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    document_id = db.Column(db.Integer, db.ForeignKey('document.id'), nullable=True)  # Optional: chat about specific document
    title = db.Column(db.String(256), default='New Chat')
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    updated_at = db.Column(db.DateTime, default=datetime.utcnow, onupdate=datetime.utcnow)
    
    # Relationships
    user = db.relationship('User', backref='chat_sessions')
    document = db.relationship('Document', backref='chat_sessions')
    messages = db.relationship('ChatMessage', backref='session', lazy='dynamic', cascade='all, delete-orphan', order_by='ChatMessage.created_at')
    
    def __repr__(self):
        return f'<ChatSession {self.id} - {self.title}>'
    
    def message_count(self):
        """Return the number of messages in this session."""
        return self.messages.count()


class ChatMessage(db.Model):
    """Model for individual messages in a chat session."""
    id = db.Column(db.Integer, primary_key=True)
    session_id = db.Column(db.Integer, db.ForeignKey('chat_session.id'), nullable=False)
    role = db.Column(db.String(16), nullable=False)  # 'user' or 'assistant'
    content = db.Column(db.Text, nullable=False)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    
    def __repr__(self):
        return f'<ChatMessage {self.id} - {self.role}>'


@login_manager.user_loader
def load_user(user_id):
    return User.query.get(int(user_id))


def allowed_file(filename: str) -> bool:
    # allow common document/image types; be permissive
    return '.' in filename


def human_year_label(y: int) -> str:
    mapping = {1: '1st Year', 2: '2nd Year', 3: '3rd Year', 4: '4th Year'}
    return mapping.get(y, f'{y} Year')


def get_subject_color_class(subject: str) -> str:
    """Get the color class for a subject."""
    subject_lower = subject.lower()
    
    color_mapping = {
        'math': 'math',
        'mathematics': 'math',
        'science': 'science',
        'english': 'english',
        'history': 'history',
        'physics': 'physics',
        'chemistry': 'chemistry',
        'biology': 'biology',
        'geography': 'geography',
        'computer': 'computer',
        'cs': 'computer',
        'programming': 'computer',
    }
    
    for key, value in color_mapping.items():
        if key in subject_lower:
            return value
    
    return 'default'


def can_access_document(user, document):
    """Check if user can access a document (owner or has share permission)."""
    if document.user_id == user.id:
        return True
    
    # Check if document is shared with user
    share = SharePermission.query.filter_by(
        shared_with_id=user.id,
        document_id=document.id
    ).first()
    
    return share is not None


def can_edit_document(user, document):
    """Check if user can edit a document (owner or has editor/admin permission)."""
    if document.user_id == user.id:
        return True
    
    # Check if user has editor or admin permission
    share = SharePermission.query.filter_by(
        shared_with_id=user.id,
        document_id=document.id
    ).filter(SharePermission.permission.in_(['editor', 'admin'])).first()
    
    return share is not None


def can_access_collection(user, collection):
    """Check if user can access a collection (owner or has share permission)."""
    if collection.user_id == user.id:
        return True
    
    # Check if collection is shared with user
    share = SharePermission.query.filter_by(
        shared_with_id=user.id,
        collection_id=collection.id
    ).first()
    
    return share is not None


def can_edit_collection(user, collection):
    """Check if user can edit a collection (owner or has editor/admin permission)."""
    if collection.user_id == user.id:
        return True
    
    # Check if user has editor or admin permission
    share = SharePermission.query.filter_by(
        shared_with_id=user.id,
        collection_id=collection.id
    ).filter(SharePermission.permission.in_(['editor', 'admin'])).first()
    
    return share is not None


def create_notification(user_id, title, message, notification_type='info', link=None):
    """Helper function to create a notification for a user."""
    notification = Notification(
        user_id=user_id,
        title=title,
        message=message,
        type=notification_type,
        link=link
    )
    db.session.add(notification)
    db.session.commit()
    return notification


def generate_thumbnail(file_path: str, mimetype: str) -> str:
    """
    Generate a thumbnail for an uploaded file (image or PDF).
    Returns the thumbnail filename or None if generation fails.
    """
    try:
        thumbnail_dir = os.path.join(app.config['UPLOAD_FOLDER'], 'thumbnails')
        os.makedirs(thumbnail_dir, exist_ok=True)
        
        # Generate thumbnail filename
        filename = os.path.basename(file_path)
        name, _ = os.path.splitext(filename)
        thumbnail_filename = f'thumb_{name}.jpg'
        thumbnail_path = os.path.join(thumbnail_dir, thumbnail_filename)
        
        # Check if it's an image
        if mimetype.startswith('image/'):
            # Open and resize the image
            img = Image.open(file_path)
            # Convert RGBA to RGB if necessary (for PNG with transparency)
            if img.mode in ('RGBA', 'LA', 'P'):
                background = Image.new('RGB', img.size, (255, 255, 255))
                if img.mode == 'P':
                    img = img.convert('RGBA')
                background.paste(img, mask=img.split()[-1] if img.mode in ('RGBA', 'LA') else None)
                img = background
            elif img.mode != 'RGB':
                img = img.convert('RGB')
            
            # Create thumbnail (200x200 max, maintaining aspect ratio)
            img.thumbnail((200, 200), Image.Resampling.LANCZOS)
            img.save(thumbnail_path, 'JPEG', quality=85)
            return thumbnail_filename
            
        # Check if it's a PDF
        elif mimetype == 'application/pdf':
            # Convert first page of PDF to image
            images = convert_from_path(file_path, first_page=1, last_page=1, dpi=100)
            if images:
                img = images[0]
                # Convert to RGB if necessary
                if img.mode != 'RGB':
                    img = img.convert('RGB')
                # Create thumbnail
                img.thumbnail((200, 200), Image.Resampling.LANCZOS)
                img.save(thumbnail_path, 'JPEG', quality=85)
                return thumbnail_filename
        
        return None
    except Exception as e:
        print(f"Error generating thumbnail: {e}")
        return None


# ============================================================================
# S3 Storage Helper Functions
# ============================================================================

def upload_to_s3(file_path: str, s3_key: str, mimetype: str = None) -> bool:
    """
    Upload a file to S3 bucket.
    Returns True if successful, False otherwise.
    """
    if not s3_client or app.config['STORAGE_TYPE'] != 's3':
        return False
    
    try:
        extra_args = {}
        if mimetype:
            extra_args['ContentType'] = mimetype
        
        with open(file_path, 'rb') as file_data:
            s3_client.upload_fileobj(
                file_data,
                app.config['S3_BUCKET_NAME'],
                s3_key,
                ExtraArgs=extra_args
            )
        print(f"✓ Uploaded to S3: {s3_key}")
        return True
    except ClientError as e:
        print(f"✗ S3 upload failed: {e}")
        return False
    except Exception as e:
        print(f"✗ Unexpected error during S3 upload: {e}")
        return False


def download_from_s3(s3_key: str, local_path: str) -> bool:
    """
    Download a file from S3 to local filesystem.
    Returns True if successful, False otherwise.
    """
    if not s3_client:
        return False
    
    try:
        s3_client.download_file(
            app.config['S3_BUCKET_NAME'],
            s3_key,
            local_path
        )
        print(f"✓ Downloaded from S3: {s3_key}")
        return True
    except ClientError as e:
        print(f"✗ S3 download failed: {e}")
        return False
    except Exception as e:
        print(f"✗ Unexpected error during S3 download: {e}")
        return False


def delete_from_s3(s3_key: str) -> bool:
    """
    Delete a file from S3 bucket.
    Returns True if successful, False otherwise.
    """
    if not s3_client:
        return False
    
    try:
        s3_client.delete_object(
            Bucket=app.config['S3_BUCKET_NAME'],
            Key=s3_key
        )
        print(f"✓ Deleted from S3: {s3_key}")
        return True
    except ClientError as e:
        print(f"✗ S3 deletion failed: {e}")
        return False
    except Exception as e:
        print(f"✗ Unexpected error during S3 deletion: {e}")
        return False


def generate_presigned_url(s3_key: str, expiration: int = 3600, as_attachment: bool = False, download_name: str = None) -> str:
    """
    Generate a presigned URL for accessing an S3 object.
    Returns the URL string or None if failed.
    
    Args:
        s3_key: The S3 object key
        expiration: URL expiration time in seconds (default 1 hour)
        as_attachment: Whether to force download
        download_name: Custom filename for download
    """
    if not s3_client:
        return None
    
    try:
        params = {
            'Bucket': app.config['S3_BUCKET_NAME'],
            'Key': s3_key
        }
        
        # Add response headers for download
        if as_attachment or download_name:
            response_disposition = f'attachment; filename="{download_name}"' if download_name else 'attachment'
            params['ResponseContentDisposition'] = response_disposition
        
        url = s3_client.generate_presigned_url(
            'get_object',
            Params=params,
            ExpiresIn=expiration
        )
        return url
    except ClientError as e:
        print(f"✗ Failed to generate presigned URL: {e}")
        return None
    except Exception as e:
        print(f"✗ Unexpected error generating presigned URL: {e}")
        return None


# ============================================================================
# Azure Blob Storage Helper Functions
# ============================================================================

def upload_to_azure(file_path: str, blob_name: str, mimetype: str = None) -> bool:
    """
    Upload a file to Azure Blob Storage.
    Returns True if successful, False otherwise.
    """
    if not blob_service_client or app.config['STORAGE_TYPE'] != 'azure':
        return False
    
    try:
        blob_client = blob_service_client.get_blob_client(
            container=app.config['AZURE_CONTAINER_NAME'],
            blob=blob_name
        )
        
        content_settings = None
        if mimetype:
            from azure.storage.blob import ContentSettings
            content_settings = ContentSettings(content_type=mimetype)
        
        with open(file_path, 'rb') as data:
            blob_client.upload_blob(data, overwrite=True, content_settings=content_settings)
        
        print(f"✓ Uploaded to Azure: {blob_name}")
        return True
    except AzureError as e:
        print(f"✗ Azure upload failed: {e}")
        return False
    except Exception as e:
        print(f"✗ Unexpected error during Azure upload: {e}")
        return False


def download_from_azure(blob_name: str, local_path: str) -> bool:
    """
    Download a file from Azure Blob Storage to local filesystem.
    Returns True if successful, False otherwise.
    """
    if not blob_service_client:
        return False
    
    try:
        blob_client = blob_service_client.get_blob_client(
            container=app.config['AZURE_CONTAINER_NAME'],
            blob=blob_name
        )
        
        with open(local_path, 'wb') as download_file:
            download_file.write(blob_client.download_blob().readall())
        
        print(f"✓ Downloaded from Azure: {blob_name}")
        return True
    except AzureError as e:
        print(f"✗ Azure download failed: {e}")
        return False
    except Exception as e:
        print(f"✗ Unexpected error during Azure download: {e}")
        return False


def delete_from_azure(blob_name: str) -> bool:
    """
    Delete a file from Azure Blob Storage.
    Returns True if successful, False otherwise.
    """
    if not blob_service_client:
        return False
    
    try:
        blob_client = blob_service_client.get_blob_client(
            container=app.config['AZURE_CONTAINER_NAME'],
            blob=blob_name
        )
        blob_client.delete_blob()
        
        print(f"✓ Deleted from Azure: {blob_name}")
        return True
    except AzureError as e:
        print(f"✗ Azure deletion failed: {e}")
        return False
    except Exception as e:
        print(f"✗ Unexpected error during Azure deletion: {e}")
        return False


def generate_azure_sas_url(blob_name: str, expiration: int = 3600, as_attachment: bool = False, download_name: str = None) -> str:
    """
    Generate a SAS (Shared Access Signature) URL for accessing an Azure blob.
    Returns the URL string or None if failed.
    
    Args:
        blob_name: The blob name
        expiration: URL expiration time in seconds (default 1 hour)
        as_attachment: Whether to force download
        download_name: Custom filename for download
    """
    if not blob_service_client:
        return None
    
    try:
        blob_client = blob_service_client.get_blob_client(
            container=app.config['AZURE_CONTAINER_NAME'],
            blob=blob_name
        )
        
        # Generate SAS token
        sas_token = generate_blob_sas(
            account_name=app.config['AZURE_STORAGE_ACCOUNT_NAME'],
            container_name=app.config['AZURE_CONTAINER_NAME'],
            blob_name=blob_name,
            account_key=app.config['AZURE_STORAGE_ACCOUNT_KEY'],
            permission=BlobSasPermissions(read=True),
            expiry=datetime.utcnow() + timedelta(seconds=expiration)
        )
        
        # Construct URL
        url = f"{blob_client.url}?{sas_token}"
        
        # Add content disposition for downloads
        if as_attachment and download_name:
            url += f"&response-content-disposition=attachment; filename=\"{download_name}\""
        elif as_attachment:
            url += "&response-content-disposition=attachment"
        
        return url
    except AzureError as e:
        print(f"✗ Failed to generate Azure SAS URL: {e}")
        return None
    except Exception as e:
        print(f"✗ Unexpected error generating Azure SAS URL: {e}")
        return None


# Authentication routes
@app.route('/login')
def login():
    if current_user.is_authenticated:
        return redirect(url_for('index'))
    return render_template('login.html')


@app.route('/login/google')
def login_google():
    redirect_uri = url_for('google_callback', _external=True)
    return google.authorize_redirect(redirect_uri)


@app.route('/login/google/callback')
def google_callback():
    try:
        token = google.authorize_access_token()
        user_info = token.get('userinfo')
        
        if not user_info:
            flash('Failed to get user info from Google', 'danger')
            return redirect(url_for('login'))
        
        # Check if user exists
        user = User.query.filter_by(google_id=user_info['sub']).first()
        
        if not user:
            # Create new user
            user = User(
                email=user_info['email'],
                name=user_info.get('name'),
                google_id=user_info['sub'],
                profile_pic=user_info.get('picture')
            )
            db.session.add(user)
            db.session.commit()
            flash(f'Welcome {user.name}! Your account has been created.', 'success')
        else:
            # Update existing user info
            user.name = user_info.get('name')
            user.profile_pic = user_info.get('picture')
            db.session.commit()
            flash(f'Welcome back, {user.name}!', 'success')
        
        login_user(user)
        return redirect(url_for('index'))
        
    except Exception as e:
        flash(f'Authentication failed: {str(e)}', 'danger')
        return redirect(url_for('login'))


@app.route('/logout')
@login_required
def logout():
    logout_user()
    flash('You have been logged out.', 'info')
    return redirect(url_for('login'))


# ============================================================================
# AI Helper Functions
# ============================================================================

def extract_text_from_pdf(file_path):
    """Extract text from PDF file."""
    try:
        text = ""
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
        return text.strip()
    except Exception as e:
        print(f"Error extracting text from PDF: {e}")
        return None


def extract_text_from_docx(file_path):
    """Extract text from DOCX file."""
    try:
        doc = DocxDocument(file_path)
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        return text.strip()
    except Exception as e:
        print(f"Error extracting text from DOCX: {e}")
        return None


def extract_text_from_image(file_path):
    """Extract text from image using OCR."""
    if not os.path.exists(app.config['TESSERACT_CMD']):
        print("Tesseract OCR not available")
        return None
    
    try:
        image = Image.open(file_path)
        text = pytesseract.image_to_string(image, lang=app.config['OCR_LANGUAGE'])
        return text.strip()
    except Exception as e:
        print(f"Error extracting text from image: {e}")
        return None


def extract_text_from_document(file_path, mimetype):
    """Extract text from document based on file type."""
    if mimetype == 'application/pdf':
        return extract_text_from_pdf(file_path)
    elif mimetype == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
        return extract_text_from_docx(file_path)
    elif mimetype.startswith('image/'):
        return extract_text_from_image(file_path)
    return None


def generate_summary(text, max_length=None):
    """Generate AI summary of text using Google Gemini."""
    if not gemini_model and not openai_client:
        return None
    
    if not text:
        return None
    
    if max_length is None:
        max_length = app.config['AI_SUMMARY_MAX_LENGTH']
    
    try:
        # Truncate text if too long
        max_input_chars = 12000
        if len(text) > max_input_chars:
            text = text[:max_input_chars] + "..."
        
        # Try Gemini first (free)
        if gemini_model:
            try:
                prompt = f"Summarize the following academic document in {max_length} characters or less. Be concise and focus on key points:\n\n{text}"
                response = gemini_model.generate_content(prompt)
                summary = response.text.strip()
                return summary[:max_length]
            except Exception as e:
                print(f"Gemini summary error: {e}")
                # Fall through to OpenAI if available
        
        # Fallback to OpenAI if Gemini fails
        if openai_client:
            response = openai_client.chat.completions.create(
                model="gpt-3.5-turbo",
                messages=[
                    {"role": "system", "content": "You are a helpful assistant that summarizes academic documents concisely."},
                    {"role": "user", "content": f"Summarize the following document in {max_length} characters or less:\n\n{text}"}
                ],
                max_tokens=200,
                temperature=0.5
            )
            summary = response.choices[0].message.content.strip()
            return summary[:max_length]
        
        return None
    except Exception as e:
        print(f"Error generating summary: {e}")
        return None


def generate_smart_tags(text, subject=None):
    """Generate smart tags using Google Gemini and NLP."""
    if not text:
        return []
    
    tags = []
    
    # Method 1: Extract key terms using NLP
    try:
        # Tokenize and remove stopwords
        words = word_tokenize(text.lower())
        words = [w for w in words if w.isalnum() and w not in stop_words and len(w) > 3]
        
        # Get most common words
        from collections import Counter
        word_freq = Counter(words)
        common_words = [word for word, _ in word_freq.most_common(10)]
        tags.extend(common_words[:3])
    except:
        pass
    
    # Method 2: Use Gemini for better tags (free)
    if gemini_model and len(text) > 50:
        try:
            # Truncate text if too long
            max_input_chars = 6000
            if len(text) > max_input_chars:
                text = text[:max_input_chars] + "..."
            
            prompt = f"Extract {app.config['AI_TAGS_COUNT']} relevant keywords/tags from this academic document."
            if subject:
                prompt += f" Subject: {subject}."
            prompt += f"\n\nDocument:\n{text}\n\nProvide only the tags, comma-separated, lowercase:"
            
            response = gemini_model.generate_content(prompt)
            ai_tags = response.text.strip()
            # Parse comma-separated tags
            ai_tags_list = [tag.strip().lower() for tag in ai_tags.split(',') if tag.strip()]
            tags.extend(ai_tags_list)
        except Exception as e:
            print(f"Gemini tags error: {e}")
            # Try OpenAI fallback
            if openai_client:
                try:
                    prompt = f"Extract {app.config['AI_TAGS_COUNT']} relevant keywords/tags from this academic document."
                    if subject:
                        prompt += f" Subject: {subject}."
                    prompt += f"\n\nDocument:\n{text}\n\nProvide only the tags, comma-separated:"
                    
                    response = openai_client.chat.completions.create(
                        model="gpt-3.5-turbo",
                        messages=[
                            {"role": "system", "content": "You are a helpful assistant that extracts relevant keywords from academic content."},
                            {"role": "user", "content": prompt}
                        ],
                        max_tokens=50,
                        temperature=0.3
                    )
                    
                    ai_tags = response.choices[0].message.content.strip()
                    ai_tags_list = [tag.strip() for tag in ai_tags.split(',') if tag.strip()]
                    tags.extend(ai_tags_list)
                except Exception as e2:
                    print(f"Error generating AI tags: {e2}")
    
    # Remove duplicates and limit count
    tags = list(dict.fromkeys(tags))  # Preserve order while removing duplicates
    return tags[:app.config['AI_TAGS_COUNT']]


def analyze_document(document_id):
    """Run full AI analysis on a document: extract text, generate summary, generate tags."""
    document = Document.query.get(document_id)
    if not document:
        return False
    
    # Get file path
    if document.storage_type == 'local':
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], document.stored_filename)
        if not os.path.exists(file_path):
            return False
    else:
        # For S3/Azure, would need to download file first
        return False
    
    # Extract text
    extracted_text = extract_text_from_document(file_path, document.mimetype)
    if extracted_text:
        document.extracted_text = extracted_text
        
        # Generate summary
        summary = generate_summary(extracted_text)
        if summary:
            document.summary = summary
        
        # Generate smart tags
        smart_tags = generate_smart_tags(extracted_text, document.subject)
        if smart_tags:
            document.ai_tags = ', '.join(smart_tags)
        
        # Update timestamp
        document.last_analyzed = datetime.utcnow()
        
        db.session.commit()
        return True
    
    return False


def search_documents_fulltext(query, user_id, limit=None):
    """Search documents by full-text search in extracted text and summaries."""
    if limit is None:
        limit = app.config['SEARCH_RESULTS_LIMIT']
    
    # Search in extracted_text, summary, original_filename, subject, tags
    search_pattern = f"%{query}%"
    
    results = Document.query.filter(
        Document.user_id == user_id,
        db.or_(
            Document.extracted_text.ilike(search_pattern),
            Document.summary.ilike(search_pattern),
            Document.original_filename.ilike(search_pattern),
            Document.subject.ilike(search_pattern),
            Document.tags.ilike(search_pattern),
            Document.ai_tags.ilike(search_pattern)
        )
    ).limit(limit).all()
    
    return results


def get_document_recommendations(document_id, count=None):
    """Get recommended documents based on content similarity."""
    if count is None:
        count = app.config['RECOMMENDATIONS_COUNT']
    
    document = Document.query.get(document_id)
    if not document or not document.extracted_text:
        return []
    
    # Get all documents from same user with extracted text
    all_docs = Document.query.filter(
        Document.user_id == document.user_id,
        Document.id != document_id,
        Document.extracted_text.isnot(None)
    ).all()
    
    if not all_docs:
        return []
    
    try:
        # Create TF-IDF vectors
        vectorizer = TfidfVectorizer(max_features=100, stop_words='english')
        
        # Prepare texts
        texts = [document.extracted_text] + [doc.extracted_text for doc in all_docs]
        tfidf_matrix = vectorizer.fit_transform(texts)
        
        # Calculate cosine similarity
        similarities = cosine_similarity(tfidf_matrix[0:1], tfidf_matrix[1:]).flatten()
        
        # Get top recommendations
        top_indices = similarities.argsort()[-count:][::-1]
        recommendations = [all_docs[i] for i in top_indices if similarities[i] > 0.1]
        
        return recommendations
    except Exception as e:
        print(f"Error getting recommendations: {e}")
        return []


@app.route('/')
@login_required
def index():
    years = [r[0] for r in db.session.query(Document.year).filter_by(user_id=current_user.id).distinct().order_by(Document.year).all()]
    years = sorted([y for y in years if y is not None])
    
    # Get total document count for current user
    total_docs = Document.query.filter_by(user_id=current_user.id).count()
    
    return render_template('index.html', years=years, human_year_label=human_year_label, total_docs=total_docs)


@app.route('/upload', methods=['GET', 'POST'])
@login_required
def upload():
    if request.method == 'POST':
        file = request.files.get('file')
        year = request.form.get('year')
        subject = request.form.get('subject')
        tags = request.form.get('tags')

        if not file or file.filename == '':
            flash('No file selected', 'warning')
            return redirect(request.url)

        try:
            year_int = int(year)
        except Exception:
            flash('Year must be a number (e.g., 1, 2, 3, 4)', 'warning')
            return redirect(request.url)

        if not subject:
            flash('Subject is required', 'warning')
            return redirect(request.url)

        if file and allowed_file(file.filename):
            orig = secure_filename(file.filename)
            ext = os.path.splitext(orig)[1]
            stored = f"{uuid4().hex}{ext}"
            save_path = os.path.join(app.config['UPLOAD_FOLDER'], stored)
            file.save(save_path)

            # Get file size
            file_size = os.path.getsize(save_path)
            
            # Generate thumbnail for images and PDFs
            thumbnail = generate_thumbnail(save_path, file.mimetype)
            
            # Upload to cloud storage if configured
            storage_type = app.config['STORAGE_TYPE']
            
            if storage_type == 's3':
                # Upload thumbnail to S3
                if thumbnail:
                    thumbnail_path = os.path.join(app.config['UPLOAD_FOLDER'], 'thumbnails', thumbnail)
                    s3_thumbnail_key = f"thumbnails/{thumbnail}"
                    upload_to_s3(thumbnail_path, s3_thumbnail_key, 'image/jpeg')
                
                # Upload main file to S3
                s3_key = f"documents/{stored}"
                success = upload_to_s3(save_path, s3_key, file.mimetype)
                if success:
                    # Remove local file after successful S3 upload
                    os.remove(save_path)
                    if thumbnail:
                        # Also remove local thumbnail
                        thumbnail_path = os.path.join(app.config['UPLOAD_FOLDER'], 'thumbnails', thumbnail)
                        if os.path.exists(thumbnail_path):
                            os.remove(thumbnail_path)
                else:
                    flash('Failed to upload to cloud storage', 'danger')
                    storage_type = 'local'  # Fall back to local
                    
            elif storage_type == 'azure':
                # Upload thumbnail to Azure
                if thumbnail:
                    thumbnail_path = os.path.join(app.config['UPLOAD_FOLDER'], 'thumbnails', thumbnail)
                    azure_thumbnail_blob = f"thumbnails/{thumbnail}"
                    upload_to_azure(thumbnail_path, azure_thumbnail_blob, 'image/jpeg')
                
                # Upload main file to Azure
                azure_blob_name = f"documents/{stored}"
                success = upload_to_azure(save_path, azure_blob_name, file.mimetype)
                if success:
                    # Remove local file after successful Azure upload
                    os.remove(save_path)
                    if thumbnail:
                        # Also remove local thumbnail
                        thumbnail_path = os.path.join(app.config['UPLOAD_FOLDER'], 'thumbnails', thumbnail)
                        if os.path.exists(thumbnail_path):
                            os.remove(thumbnail_path)
                else:
                    flash('Failed to upload to cloud storage', 'danger')
                    storage_type = 'local'  # Fall back to local

            doc = Document(
                original_filename=orig,
                stored_filename=stored,
                year=year_int,
                subject=subject,
                mimetype=file.mimetype,
                size=file_size,
                thumbnail_filename=thumbnail,
                storage_type=storage_type,
                user_id=current_user.id
            )
            # Set tags using the new tag system
            doc.set_tags_from_string(tags)
            db.session.add(doc)
            db.session.commit()
            flash('File uploaded successfully', 'success')
            return redirect(url_for('year_view', year=year_int))

    # GET -> show form
    return render_template('upload.html')


@app.route('/upload-multiple', methods=['POST'])
@login_required
def upload_multiple():
    """Handle multiple file uploads via AJAX and return JSON response."""
    files = request.files.getlist('files[]')
    year = request.form.get('year')
    subject = request.form.get('subject')
    tags = request.form.get('tags')
    
    # Validation
    if not files or len(files) == 0:
        return jsonify({'success': False, 'error': 'No files selected'}), 400
    
    try:
        year_int = int(year)
    except (ValueError, TypeError):
        return jsonify({'success': False, 'error': 'Invalid year value'}), 400
    
    if not subject:
        return jsonify({'success': False, 'error': 'Subject is required'}), 400
    
    results = []
    success_count = 0
    
    for file in files:
        result = {
            'filename': file.filename,
            'success': False,
            'error': None
        }
        
        try:
            if not file or file.filename == '':
                result['error'] = 'Empty file'
                results.append(result)
                continue
            
            if not allowed_file(file.filename):
                result['error'] = 'File type not allowed'
                results.append(result)
                continue
            
            # Process the file
            orig = secure_filename(file.filename)
            ext = os.path.splitext(orig)[1]
            stored = f"{uuid4().hex}{ext}"
            save_path = os.path.join(app.config['UPLOAD_FOLDER'], stored)
            file.save(save_path)
            
            # Get file size
            file_size = os.path.getsize(save_path)
            
            # Generate thumbnail for images and PDFs
            thumbnail = generate_thumbnail(save_path, file.mimetype)
            
            # Upload to cloud storage if configured
            storage_type = app.config['STORAGE_TYPE']
            
            if storage_type == 's3':
                # Upload thumbnail to S3
                if thumbnail:
                    thumbnail_path = os.path.join(app.config['UPLOAD_FOLDER'], 'thumbnails', thumbnail)
                    s3_thumbnail_key = f"thumbnails/{thumbnail}"
                    upload_to_s3(thumbnail_path, s3_thumbnail_key, 'image/jpeg')
                
                # Upload main file to S3
                s3_key = f"documents/{stored}"
                success = upload_to_s3(save_path, s3_key, file.mimetype)
                if success:
                    os.remove(save_path)
                    if thumbnail:
                        thumbnail_path = os.path.join(app.config['UPLOAD_FOLDER'], 'thumbnails', thumbnail)
                        if os.path.exists(thumbnail_path):
                            os.remove(thumbnail_path)
                else:
                    storage_type = 'local'
                    
            elif storage_type == 'azure':
                # Upload thumbnail to Azure
                if thumbnail:
                    thumbnail_path = os.path.join(app.config['UPLOAD_FOLDER'], 'thumbnails', thumbnail)
                    azure_thumbnail_blob = f"thumbnails/{thumbnail}"
                    upload_to_azure(thumbnail_path, azure_thumbnail_blob, 'image/jpeg')
                
                # Upload main file to Azure
                azure_blob_name = f"documents/{stored}"
                success = upload_to_azure(save_path, azure_blob_name, file.mimetype)
                if success:
                    os.remove(save_path)
                    if thumbnail:
                        thumbnail_path = os.path.join(app.config['UPLOAD_FOLDER'], 'thumbnails', thumbnail)
                        if os.path.exists(thumbnail_path):
                            os.remove(thumbnail_path)
                else:
                    storage_type = 'local'
            
            # Save to database
            doc = Document(
                original_filename=orig,
                stored_filename=stored,
                year=year_int,
                subject=subject,
                mimetype=file.mimetype,
                size=file_size,
                thumbnail_filename=thumbnail,
                storage_type=storage_type,
                user_id=current_user.id
            )
            doc.set_tags_from_string(tags)
            db.session.add(doc)
            db.session.commit()
            
            result['success'] = True
            success_count += 1
            
        except Exception as e:
            result['error'] = str(e)
            print(f"Error uploading {file.filename}: {e}")
        
        results.append(result)
    
    return jsonify({
        'success': True,
        'total': len(files),
        'successful': success_count,
        'failed': len(files) - success_count,
        'results': results
    })


@app.route('/create-note', methods=['GET', 'POST'])
@login_required
def create_note():
    if request.method == 'POST':
        title = request.form.get('title')
        content = request.form.get('content')
        year = request.form.get('year')
        subject = request.form.get('subject')
        tags = request.form.get('tags')

        if not title or not title.strip():
            flash('Note title is required', 'warning')
            return redirect(request.url)

        if not content or not content.strip():
            flash('Note content is required', 'warning')
            return redirect(request.url)

        try:
            year_int = int(year)
        except Exception:
            flash('Year must be a number (e.g., 1, 2, 3, 4)', 'warning')
            return redirect(request.url)

        if not subject:
            flash('Subject is required', 'warning')
            return redirect(request.url)

        # Create a text file with the note content
        safe_title = secure_filename(title)
        if not safe_title.endswith('.txt'):
            safe_title += '.txt'
        
        stored = f"{uuid4().hex}.txt"
        save_path = os.path.join(app.config['UPLOAD_FOLDER'], stored)
        
        with open(save_path, 'w', encoding='utf-8') as f:
            f.write(content)

        doc = Document(
            original_filename=safe_title,
            stored_filename=stored,
            year=year_int,
            subject=subject,
            mimetype='text/plain',
            size=os.path.getsize(save_path),
            user_id=current_user.id
        )
        # Set tags using the new tag system
        doc.set_tags_from_string(tags)
        db.session.add(doc)
        db.session.commit()
        flash('Note created successfully', 'success')
        return redirect(url_for('year_view', year=year_int))

    # GET -> show form
    return render_template('create_note.html')


@app.route('/year/<int:year>')
@login_required
def year_view(year: int):
    subject = request.args.get('subject', type=str)
    tags = request.args.get('tags', type=str)
    page = request.args.get('page', 1, type=int)
    per_page = 10

    q = Document.query.filter_by(year=year, user_id=current_user.id)
    if subject:
        q = q.filter(Document.subject.ilike(f'%{subject}%'))
    if tags:
        # simple tags search: every provided tag must appear in stored tags string
        for t in [t.strip() for t in tags.split(',') if t.strip()]:
            q = q.filter(Document.tags.ilike(f'%{t}%'))

    pagination = q.order_by(Document.upload_date.desc()).paginate(
        page=page, per_page=per_page, error_out=False
    )
    documents = pagination.items
    
    # collect distinct subjects for a quick filter UI
    subjects = [r[0] for r in db.session.query(Document.subject).filter_by(year=year, user_id=current_user.id).distinct().order_by(Document.subject).all()]
    
    return render_template(
        'year.html', 
        year=year, 
        documents=documents, 
        pagination=pagination,
        subjects=subjects, 
        human_year_label=human_year_label,
        get_subject_color_class=get_subject_color_class
    )


@app.route('/search')
@login_required
def search():
    query = request.args.get('q', '').strip()
    page = request.args.get('page', 1, type=int)
    per_page = 10
    
    if not query:
        flash('Please enter a search term', 'warning')
        return redirect(url_for('index'))
    
    # Search across multiple fields
    search_term = f'%{query}%'
    q = Document.query.filter_by(user_id=current_user.id).filter(
        db.or_(
            Document.original_filename.ilike(search_term),
            Document.subject.ilike(search_term),
            Document.tags.ilike(search_term)
        )
    )
    
    # For text files, also search content
    text_docs = Document.query.filter_by(user_id=current_user.id, mimetype='text/plain').all()
    matching_text_doc_ids = []
    
    for doc in text_docs:
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], doc.stored_filename)
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()
                if query.lower() in content.lower():
                    matching_text_doc_ids.append(doc.id)
        except Exception:
            pass
    
    # Combine results
    if matching_text_doc_ids:
        q = q.union(
            Document.query.filter(Document.id.in_(matching_text_doc_ids))
        )
    
    pagination = q.order_by(Document.upload_date.desc()).paginate(
        page=page, per_page=per_page, error_out=False
    )
    documents = pagination.items
    
    return render_template(
        'search.html',
        query=query,
        documents=documents,
        pagination=pagination,
        human_year_label=human_year_label,
        get_subject_color_class=get_subject_color_class
    )


@app.route('/tags')
@login_required
def tags_list():
    """Show all tags with document counts."""
    # Get all tags that have documents belonging to current user
    tags_with_counts = db.session.query(
        Tag.id, Tag.name, Tag.slug, db.func.count(Document.id).label('count')
    ).select_from(Tag).join(document_tags, Tag.id == document_tags.c.tag_id).join(
        Document, Document.id == document_tags.c.document_id
    ).filter(
        Document.user_id == current_user.id
    ).group_by(Tag.id, Tag.name, Tag.slug).order_by(db.desc('count')).all()
    
    return render_template('tags.html', tags=tags_with_counts)


@app.route('/tag/<slug>')
@login_required
def tag_view(slug):
    """Show all documents with a specific tag."""
    tag = Tag.query.filter_by(slug=slug).first_or_404()
    page = request.args.get('page', 1, type=int)
    per_page = 10
    
    # Get documents with this tag that belong to current user
    q = Document.query.join(document_tags).join(Tag).filter(
        Tag.slug == slug,
        Document.user_id == current_user.id
    )
    
    pagination = q.order_by(Document.upload_date.desc()).paginate(
        page=page, per_page=per_page, error_out=False
    )
    documents = pagination.items
    
    return render_template(
        'tag.html',
        tag=tag,
        documents=documents,
        pagination=pagination,
        human_year_label=human_year_label,
        get_subject_color_class=get_subject_color_class
    )


@app.route('/download/<int:doc_id>')
@login_required
def download(doc_id: int):
    doc = Document.query.filter_by(id=doc_id, user_id=current_user.id).first_or_404()
    
    # If file is in S3, generate presigned URL and redirect
    if doc.storage_type == 's3':
        s3_key = f"documents/{doc.stored_filename}"
        url = generate_presigned_url(s3_key, expiration=300, as_attachment=True, download_name=doc.original_filename)
        if url:
            return redirect(url)
        else:
            flash('Failed to generate download link', 'danger')
            return redirect(request.referrer or url_for('index'))
    
    # If file is in Azure, generate SAS URL and redirect
    elif doc.storage_type == 'azure':
        azure_blob_name = f"documents/{doc.stored_filename}"
        url = generate_azure_sas_url(azure_blob_name, expiration=300, as_attachment=True, download_name=doc.original_filename)
        if url:
            return redirect(url)
        else:
            flash('Failed to generate download link', 'danger')
            return redirect(request.referrer or url_for('index'))
    
    # Otherwise serve from local storage
    return send_from_directory(app.config['UPLOAD_FOLDER'], doc.stored_filename, as_attachment=True, download_name=doc.original_filename)


@app.route('/preview/<int:doc_id>')
@login_required
def preview(doc_id: int):
    doc = Document.query.filter_by(id=doc_id, user_id=current_user.id).first_or_404()
    previewable = False
    view_type = 'other'
    text_content = None
    file_url = None
    
    if doc.mimetype and doc.mimetype.startswith('image/'):
        previewable = True
        view_type = 'image'
    elif doc.mimetype == 'application/pdf' or (doc.original_filename.lower().endswith('.pdf')):
        previewable = True
        view_type = 'pdf'
    elif doc.mimetype == 'text/plain' or (doc.original_filename.lower().endswith('.txt')):
        previewable = True
        view_type = 'text'
        # Read the text file content
        if doc.storage_type == 's3':
            # Download from S3 to temporary location
            import tempfile
            temp_file = tempfile.NamedTemporaryFile(mode='w+', delete=False, suffix='.txt')
            s3_key = f"documents/{doc.stored_filename}"
            if download_from_s3(s3_key, temp_file.name):
                try:
                    with open(temp_file.name, 'r', encoding='utf-8') as f:
                        text_content = f.read()
                    os.remove(temp_file.name)
                except Exception as e:
                    text_content = f"Error reading file: {str(e)}"
            else:
                text_content = "Error downloading file from cloud storage"
        elif doc.storage_type == 'azure':
            # Download from Azure to temporary location
            import tempfile
            temp_file = tempfile.NamedTemporaryFile(mode='w+', delete=False, suffix='.txt')
            azure_blob_name = f"documents/{doc.stored_filename}"
            if download_from_azure(azure_blob_name, temp_file.name):
                try:
                    with open(temp_file.name, 'r', encoding='utf-8') as f:
                        text_content = f.read()
                    os.remove(temp_file.name)
                except Exception as e:
                    text_content = f"Error reading file: {str(e)}"
            else:
                text_content = "Error downloading file from cloud storage"
        else:
            file_path = os.path.join(app.config['UPLOAD_FOLDER'], doc.stored_filename)
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    text_content = f.read()
            except Exception as e:
                text_content = f"Error reading file: {str(e)}"
    
    # Generate file URL for preview
    if doc.storage_type == 's3' and view_type in ['image', 'pdf']:
        s3_key = f"documents/{doc.stored_filename}"
        file_url = generate_presigned_url(s3_key, expiration=3600)
    elif doc.storage_type == 'azure' and view_type in ['image', 'pdf']:
        azure_blob_name = f"documents/{doc.stored_filename}"
        file_url = generate_azure_sas_url(azure_blob_name, expiration=3600)
    else:
        file_url = url_for('uploaded_file', filename=doc.stored_filename)

    return render_template('preview.html', doc=doc, previewable=previewable, view_type=view_type, text_content=text_content, file_url=file_url)


@app.route('/edit/<int:doc_id>', methods=['GET', 'POST'])
@login_required
def edit_document(doc_id: int):
    doc = Document.query.filter_by(id=doc_id, user_id=current_user.id).first_or_404()
    
    if request.method == 'POST':
        year = request.form.get('year')
        subject = request.form.get('subject')
        tags = request.form.get('tags')
        
        try:
            year_int = int(year)
        except Exception:
            flash('Year must be a number (e.g., 1, 2, 3, 4)', 'warning')
            return redirect(request.url)
        
        if not subject:
            flash('Subject is required', 'warning')
            return redirect(request.url)
        
        # Update document metadata
        doc.year = year_int
        doc.subject = subject
        
        # Update tags using the new tag system
        doc.set_tags_from_string(tags)
        
        db.session.commit()
        flash('Document updated successfully', 'success')
        return redirect(url_for('year_view', year=year_int))
    
    # GET -> show edit form
    tags_string = ', '.join([tag.name for tag in doc.tag_objects]) if doc.tag_objects else ''
    return render_template('edit.html', doc=doc, tags_string=tags_string, human_year_label=human_year_label)


@app.route('/delete/<int:doc_id>', methods=['POST'])
@login_required
def delete_document(doc_id: int):
    doc = Document.query.filter_by(id=doc_id, user_id=current_user.id).first_or_404()
    
    # Delete the file from storage (S3, Azure, or local)
    if doc.storage_type == 's3':
        # Delete from S3
        s3_key = f"documents/{doc.stored_filename}"
        delete_from_s3(s3_key)
        
        # Delete thumbnail from S3 if exists
        if doc.thumbnail_filename:
            s3_thumbnail_key = f"thumbnails/{doc.thumbnail_filename}"
            delete_from_s3(s3_thumbnail_key)
            
    elif doc.storage_type == 'azure':
        # Delete from Azure
        azure_blob_name = f"documents/{doc.stored_filename}"
        delete_from_azure(azure_blob_name)
        
        # Delete thumbnail from Azure if exists
        if doc.thumbnail_filename:
            azure_thumbnail_blob = f"thumbnails/{doc.thumbnail_filename}"
            delete_from_azure(azure_thumbnail_blob)
            
    else:
        # Delete from local storage
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], doc.stored_filename)
        try:
            if os.path.exists(file_path):
                os.remove(file_path)
        except Exception as e:
            flash(f'Error deleting file: {str(e)}', 'danger')
            return redirect(request.referrer or url_for('index'))
        
        # Delete thumbnail from local storage if exists
        if doc.thumbnail_filename:
            thumbnail_path = os.path.join(app.config['UPLOAD_FOLDER'], 'thumbnails', doc.thumbnail_filename)
            try:
                if os.path.exists(thumbnail_path):
                    os.remove(thumbnail_path)
            except Exception as e:
                print(f"Error deleting thumbnail: {e}")
    
    # Delete from database
    db.session.delete(doc)
    db.session.commit()
    
    flash('Document deleted successfully', 'success')
    return redirect(request.referrer or url_for('index'))


# ===== SHARING & COLLABORATION ROUTES =====

@app.route('/share/document/<int:doc_id>', methods=['GET', 'POST'])
@login_required
def share_document(doc_id):
    """Share a document with another user."""
    doc = Document.query.filter_by(id=doc_id, user_id=current_user.id).first_or_404()
    
    if request.method == 'POST':
        email = request.form.get('email', '').strip()
        permission = request.form.get('permission', 'viewer')
        
        if not email:
            flash('Please provide an email address', 'danger')
            return redirect(url_for('share_document', doc_id=doc_id))
        
        # Find the user by email
        user = User.query.filter_by(email=email).first()
        if not user:
            flash(f'No user found with email: {email}', 'danger')
            return redirect(url_for('share_document', doc_id=doc_id))
        
        if user.id == current_user.id:
            flash('You cannot share with yourself', 'warning')
            return redirect(url_for('share_document', doc_id=doc_id))
        
        # Check if already shared
        existing = SharePermission.query.filter_by(
            shared_by_id=current_user.id,
            shared_with_id=user.id,
            document_id=doc_id
        ).first()
        
        if existing:
            # Update permission
            existing.permission = permission
            db.session.commit()
            flash(f'Updated sharing permission for {user.name or user.email}', 'success')
        else:
            # Create new share
            share = SharePermission(
                shared_by_id=current_user.id,
                shared_with_id=user.id,
                document_id=doc_id,
                permission=permission
            )
            db.session.add(share)
            db.session.commit()
            
            # Create notification
            create_notification(
                user_id=user.id,
                title='Document Shared',
                message=f'{current_user.name or current_user.email} shared "{doc.original_filename}" with you',
                notification_type='share',
                link=url_for('preview', doc_id=doc_id)
            )
            
            flash(f'Document shared with {user.name or user.email}', 'success')
        
        return redirect(url_for('share_document', doc_id=doc_id))
    
    # GET request - show sharing interface
    shares = SharePermission.query.filter_by(
        shared_by_id=current_user.id,
        document_id=doc_id
    ).all()
    
    return render_template('share_document.html', doc=doc, shares=shares)


@app.route('/share/collection/<int:collection_id>', methods=['GET', 'POST'])
@login_required
def share_collection(collection_id):
    """Share a collection with another user."""
    collection = Collection.query.filter_by(id=collection_id, user_id=current_user.id).first_or_404()
    
    if request.method == 'POST':
        email = request.form.get('email', '').strip()
        permission = request.form.get('permission', 'viewer')
        
        if not email:
            flash('Please provide an email address', 'danger')
            return redirect(url_for('share_collection', collection_id=collection_id))
        
        # Find the user by email
        user = User.query.filter_by(email=email).first()
        if not user:
            flash(f'No user found with email: {email}', 'danger')
            return redirect(url_for('share_collection', collection_id=collection_id))
        
        if user.id == current_user.id:
            flash('You cannot share with yourself', 'warning')
            return redirect(url_for('share_collection', collection_id=collection_id))
        
        # Check if already shared
        existing = SharePermission.query.filter_by(
            shared_by_id=current_user.id,
            shared_with_id=user.id,
            collection_id=collection_id
        ).first()
        
        if existing:
            # Update permission
            existing.permission = permission
            db.session.commit()
            flash(f'Updated sharing permission for {user.name or user.email}', 'success')
        else:
            # Create new share
            share = SharePermission(
                shared_by_id=current_user.id,
                shared_with_id=user.id,
                collection_id=collection_id,
                permission=permission
            )
            db.session.add(share)
            db.session.commit()
            
            # Create notification
            create_notification(
                user_id=user.id,
                title='Collection Shared',
                message=f'{current_user.name or current_user.email} shared collection "{collection.name}" with you',
                notification_type='share',
                link=url_for('collection_view', collection_id=collection_id)
            )
            
            flash(f'Collection shared with {user.name or user.email}', 'success')
        
        return redirect(url_for('share_collection', collection_id=collection_id))
    
    # GET request - show sharing interface
    shares = SharePermission.query.filter_by(
        shared_by_id=current_user.id,
        collection_id=collection_id
    ).all()
    
    return render_template('share_collection.html', collection=collection, shares=shares)


@app.route('/share/revoke/<int:share_id>', methods=['POST'])
@login_required
def revoke_share(share_id):
    """Revoke a share permission."""
    share = SharePermission.query.filter_by(id=share_id, shared_by_id=current_user.id).first_or_404()
    
    db.session.delete(share)
    db.session.commit()
    
    flash('Access revoked successfully', 'success')
    return redirect(request.referrer or url_for('index'))


@app.route('/shared-with-me')
@login_required
def shared_with_me():
    """View all documents and collections shared with current user."""
    # Get shared documents
    doc_shares = SharePermission.query.filter_by(
        shared_with_id=current_user.id
    ).filter(SharePermission.document_id.isnot(None)).all()
    
    shared_docs = [share.document for share in doc_shares if share.document]
    
    # Get shared collections
    collection_shares = SharePermission.query.filter_by(
        shared_with_id=current_user.id
    ).filter(SharePermission.collection_id.isnot(None)).all()
    
    shared_collections = [share.collection for share in collection_shares if share.collection]
    
    return render_template(
        'shared_with_me.html',
        shared_docs=shared_docs,
        shared_collections=shared_collections,
        get_subject_color_class=get_subject_color_class
    )


# ===== COLLECTIONS ROUTES =====

@app.route('/collections')
@login_required
def collections():
    """List all collections for the current user."""
    user_collections = Collection.query.filter_by(user_id=current_user.id).order_by(Collection.updated_at.desc()).all()
    
    # Calculate document count for each collection
    collections_data = []
    for collection in user_collections:
        collections_data.append({
            'collection': collection,
            'document_count': collection.document_count()
        })
    
    return render_template('collections.html', collections=collections_data)


@app.route('/collection/create', methods=['GET', 'POST'])
@login_required
def create_collection():
    """Create a new collection."""
    if request.method == 'POST':
        name = request.form.get('name', '').strip()
        description = request.form.get('description', '').strip()
        color = request.form.get('color', '#667eea').strip()
        icon = request.form.get('icon', 'bi-folder').strip()
        
        if not name:
            flash('Collection name is required', 'danger')
            return redirect(url_for('create_collection'))
        
        # Create new collection
        new_collection = Collection(
            name=name,
            description=description,
            color=color,
            icon=icon,
            user_id=current_user.id
        )
        
        db.session.add(new_collection)
        db.session.commit()
        
        flash(f'Collection "{name}" created successfully!', 'success')
        return redirect(url_for('collections'))
    
    return render_template('collection_create.html')


@app.route('/collection/<int:collection_id>')
@login_required
def view_collection(collection_id):
    """View documents in a specific collection."""
    collection = Collection.query.filter_by(id=collection_id, user_id=current_user.id).first_or_404()
    documents = collection.documents.all()
    
    return render_template('collection_view.html', collection=collection, documents=documents)


@app.route('/collection/<int:collection_id>/edit', methods=['GET', 'POST'])
@login_required
def edit_collection(collection_id):
    """Edit collection details."""
    collection = Collection.query.filter_by(id=collection_id, user_id=current_user.id).first_or_404()
    
    if request.method == 'POST':
        name = request.form.get('name', '').strip()
        description = request.form.get('description', '').strip()
        color = request.form.get('color', '#667eea').strip()
        icon = request.form.get('icon', 'bi-folder').strip()
        
        if not name:
            flash('Collection name is required', 'danger')
            return redirect(url_for('edit_collection', collection_id=collection_id))
        
        collection.name = name
        collection.description = description
        collection.color = color
        collection.icon = icon
        collection.updated_at = datetime.utcnow()
        
        db.session.commit()
        
        flash(f'Collection "{name}" updated successfully!', 'success')
        return redirect(url_for('view_collection', collection_id=collection_id))
    
    return render_template('collection_create.html', collection=collection, edit_mode=True)


@app.route('/collection/<int:collection_id>/delete', methods=['POST'])
@login_required
def delete_collection(collection_id):
    """Delete a collection (not the documents)."""
    collection = Collection.query.filter_by(id=collection_id, user_id=current_user.id).first_or_404()
    
    collection_name = collection.name
    db.session.delete(collection)
    db.session.commit()
    
    flash(f'Collection "{collection_name}" deleted successfully', 'success')
    return redirect(url_for('collections'))


@app.route('/collection/<int:collection_id>/add', methods=['POST'])
@login_required
def add_to_collection(collection_id):
    """Add documents to a collection."""
    collection = Collection.query.filter_by(id=collection_id, user_id=current_user.id).first_or_404()
    
    # Get document IDs from form (supports multiple)
    doc_ids = request.form.getlist('document_ids')
    
    if not doc_ids:
        return jsonify({'success': False, 'message': 'No documents selected'}), 400
    
    added_count = 0
    for doc_id in doc_ids:
        doc = Document.query.filter_by(id=int(doc_id), user_id=current_user.id).first()
        if doc and doc not in collection.documents:
            collection.documents.append(doc)
            added_count += 1
    
    collection.updated_at = datetime.utcnow()
    db.session.commit()
    
    if request.is_json or request.headers.get('X-Requested-With') == 'XMLHttpRequest':
        return jsonify({
            'success': True,
            'message': f'Added {added_count} document(s) to "{collection.name}"',
            'added_count': added_count
        })
    
    flash(f'Added {added_count} document(s) to "{collection.name}"', 'success')
    return redirect(request.referrer or url_for('view_collection', collection_id=collection_id))


@app.route('/collection/<int:collection_id>/remove/<int:doc_id>', methods=['POST'])
@login_required
def remove_from_collection(collection_id, doc_id):
    """Remove a document from a collection."""
    collection = Collection.query.filter_by(id=collection_id, user_id=current_user.id).first_or_404()
    doc = Document.query.filter_by(id=doc_id, user_id=current_user.id).first_or_404()
    
    if doc in collection.documents:
        collection.documents.remove(doc)
        collection.updated_at = datetime.utcnow()
        db.session.commit()
        flash(f'Document removed from "{collection.name}"', 'success')
    else:
        flash('Document not in this collection', 'warning')
    
    return redirect(url_for('view_collection', collection_id=collection_id))


@app.route('/api/documents')
@login_required
def api_get_documents():
    """API endpoint to get all user documents for collection management."""
    documents = Document.query.filter_by(user_id=current_user.id).order_by(Document.upload_date.desc()).all()
    
    documents_list = []
    for doc in documents:
        documents_list.append({
            'id': doc.id,
            'original_filename': doc.original_filename,
            'year': doc.year,
            'subject': doc.subject,
            'tags': doc.tags,
            'upload_date': doc.upload_date.strftime('%Y-%m-%d')
        })
    
    return jsonify({'documents': documents_list})


@app.route('/uploads/<path:filename>')
def uploaded_file(filename):
    # serve stored files; keep simple for a small app
    return send_from_directory(app.config['UPLOAD_FOLDER'], filename)


@app.route('/thumbnails/<path:filename>')
def thumbnail_file(filename):
    # If using S3, generate presigned URL and redirect
    if app.config['STORAGE_TYPE'] == 's3':
        s3_key = f"thumbnails/{filename}"
        url = generate_presigned_url(s3_key, expiration=3600)
        if url:
            return redirect(url)
    
    # Otherwise serve from local thumbnails directory
    thumbnail_dir = os.path.join(app.config['UPLOAD_FOLDER'], 'thumbnails')
    return send_from_directory(thumbnail_dir, filename)


# ===== COMMENTS ROUTES =====

@app.route('/document/<int:doc_id>/comments', methods=['GET', 'POST'])
@login_required
def document_comments(doc_id):
    """View and add comments to a document."""
    doc = Document.query.get_or_404(doc_id)
    
    # Check if user can access this document
    if not can_access_document(current_user, doc):
        flash('You do not have access to this document', 'danger')
        return redirect(url_for('index'))
    
    if request.method == 'POST':
        content = request.form.get('content', '').strip()
        page_number = request.form.get('page_number')
        
        if not content:
            flash('Comment cannot be empty', 'danger')
            return redirect(url_for('document_comments', doc_id=doc_id))
        
        comment = Comment(
            document_id=doc_id,
            user_id=current_user.id,
            content=content,
            page_number=int(page_number) if page_number else None
        )
        db.session.add(comment)
        db.session.commit()
        
        # Notify document owner if commenter is not the owner
        if doc.user_id != current_user.id:
            create_notification(
                user_id=doc.user_id,
                title='New Comment',
                message=f'{current_user.name or current_user.email} commented on "{doc.original_filename}"',
                notification_type='comment',
                link=url_for('document_comments', doc_id=doc_id)
            )
        
        flash('Comment added successfully', 'success')
        return redirect(url_for('document_comments', doc_id=doc_id))
    
    # GET request - show comments
    comments = Comment.query.filter_by(document_id=doc_id).order_by(Comment.created_at.desc()).all()
    
    return render_template('document_comments.html', doc=doc, comments=comments)


@app.route('/comment/<int:comment_id>/delete', methods=['POST'])
@login_required
def delete_comment(comment_id):
    """Delete a comment."""
    comment = Comment.query.get_or_404(comment_id)
    
    # Only the comment author or document owner can delete
    if comment.user_id != current_user.id and comment.document.user_id != current_user.id:
        flash('You do not have permission to delete this comment', 'danger')
        return redirect(request.referrer or url_for('index'))
    
    doc_id = comment.document_id
    db.session.delete(comment)
    db.session.commit()
    
    flash('Comment deleted successfully', 'success')
    return redirect(url_for('document_comments', doc_id=doc_id))


# ===== STUDY GROUPS ROUTES =====

@app.route('/groups')
@login_required
def study_groups():
    """View all study groups."""
    my_groups = current_user.study_groups
    created_groups = StudyGroup.query.filter_by(created_by_id=current_user.id).all()
    
    return render_template('study_groups.html', my_groups=my_groups, created_groups=created_groups)


@app.route('/group/create', methods=['GET', 'POST'])
@login_required
def create_group():
    """Create a new study group."""
    if request.method == 'POST':
        name = request.form.get('name', '').strip()
        description = request.form.get('description', '').strip()
        color = request.form.get('color', '#667eea')
        icon = request.form.get('icon', 'bi-people-fill')
        
        if not name:
            flash('Group name is required', 'danger')
            return redirect(url_for('create_group'))
        
        group = StudyGroup(
            name=name,
            description=description,
            color=color,
            icon=icon,
            created_by_id=current_user.id
        )
        
        # Add creator as first member
        group.members.append(current_user)
        
        db.session.add(group)
        db.session.commit()
        
        flash(f'Study group "{name}" created successfully', 'success')
        return redirect(url_for('group_view', group_id=group.id))
    
    return render_template('create_group.html')


@app.route('/group/<int:group_id>')
@login_required
def group_view(group_id):
    """View a study group."""
    group = StudyGroup.query.get_or_404(group_id)
    
    # Check if user is a member
    if current_user not in group.members:
        flash('You are not a member of this group', 'danger')
        return redirect(url_for('study_groups'))
    
    # Get shared documents in this group (we can extend this later)
    # For now, show documents shared by group members
    member_ids = [member.id for member in group.members]
    
    shared_docs = db.session.query(Document).join(
        SharePermission,
        Document.id == SharePermission.document_id
    ).filter(
        SharePermission.shared_with_id == current_user.id,
        SharePermission.shared_by_id.in_(member_ids)
    ).all()
    
    return render_template(
        'group_view.html',
        group=group,
        shared_docs=shared_docs,
        get_subject_color_class=get_subject_color_class
    )


@app.route('/group/<int:group_id>/invite', methods=['POST'])
@login_required
def invite_to_group(group_id):
    """Invite a user to join a study group."""
    group = StudyGroup.query.get_or_404(group_id)
    
    # Only group creator or admin can invite
    if not group.is_admin(current_user):
        flash('You do not have permission to invite members', 'danger')
        return redirect(url_for('group_view', group_id=group_id))
    
    email = request.form.get('email', '').strip()
    if not email:
        flash('Please provide an email address', 'danger')
        return redirect(url_for('group_view', group_id=group_id))
    
    user = User.query.filter_by(email=email).first()
    if not user:
        flash(f'No user found with email: {email}', 'danger')
        return redirect(url_for('group_view', group_id=group_id))
    
    if user in group.members:
        flash('User is already a member of this group', 'warning')
        return redirect(url_for('group_view', group_id=group_id))
    
    group.members.append(user)
    db.session.commit()
    
    # Create notification
    create_notification(
        user_id=user.id,
        title='Group Invitation',
        message=f'{current_user.name or current_user.email} added you to study group "{group.name}"',
        notification_type='group',
        link=url_for('group_view', group_id=group_id)
    )
    
    flash(f'{user.name or user.email} added to the group', 'success')
    return redirect(url_for('group_view', group_id=group_id))


@app.route('/group/<int:group_id>/leave', methods=['POST'])
@login_required
def leave_group(group_id):
    """Leave a study group."""
    group = StudyGroup.query.get_or_404(group_id)
    
    if current_user not in group.members:
        flash('You are not a member of this group', 'warning')
        return redirect(url_for('study_groups'))
    
    if group.created_by_id == current_user.id:
        flash('Group creator cannot leave the group. Delete the group instead.', 'warning')
        return redirect(url_for('group_view', group_id=group_id))
    
    group.members.remove(current_user)
    db.session.commit()
    
    flash(f'You left the group "{group.name}"', 'success')
    return redirect(url_for('study_groups'))


# ===== NOTIFICATIONS ROUTES =====

@app.route('/notifications')
@login_required
def notifications():
    """View all notifications."""
    notifications_list = Notification.query.filter_by(user_id=current_user.id).order_by(Notification.created_at.desc()).limit(50).all()
    
    # Mark all as read
    for notif in notifications_list:
        if not notif.read:
            notif.read = True
    db.session.commit()
    
    return render_template('notifications.html', notifications=notifications_list)


@app.route('/notifications/unread-count')
@login_required
def unread_notifications_count():
    """API endpoint to get unread notification count."""
    count = Notification.query.filter_by(user_id=current_user.id, read=False).count()
    return jsonify({'count': count})


@app.route('/notification/<int:notif_id>/mark-read', methods=['POST'])
@login_required
def mark_notification_read(notif_id):
    """Mark a notification as read."""
    notif = Notification.query.filter_by(id=notif_id, user_id=current_user.id).first_or_404()
    notif.read = True
    db.session.commit()
    return jsonify({'success': True})


# ============================================================================
# AI FEATURES ROUTES
# ============================================================================

@app.route('/document/<int:doc_id>/analyze', methods=['POST'])
@login_required
def analyze_document_route(doc_id):
    """Trigger AI analysis for a document."""
    document = Document.query.filter_by(id=doc_id, user_id=current_user.id).first_or_404()
    
    if not gemini_model and not openai_client:
        return jsonify({'success': False, 'error': 'AI features not configured'}), 400
    
    # Run analysis in the request (in production, use background task)
    success = analyze_document(doc_id)
    
    if success:
        return jsonify({
            'success': True,
            'summary': document.summary,
            'ai_tags': document.ai_tags,
            'extracted_text_length': len(document.extracted_text) if document.extracted_text else 0
        })
    else:
        return jsonify({'success': False, 'error': 'Analysis failed'}), 500


@app.route('/document/<int:doc_id>/summary')
@login_required
def get_document_summary(doc_id):
    """Get or generate document summary."""
    document = Document.query.filter_by(id=doc_id, user_id=current_user.id).first_or_404()
    
    if document.summary:
        return jsonify({'success': True, 'summary': document.summary})
    
    # Generate summary if not exists
    if document.extracted_text:
        summary = generate_summary(document.extracted_text)
        if summary:
            document.summary = summary
            db.session.commit()
            return jsonify({'success': True, 'summary': summary})
    
    return jsonify({'success': False, 'error': 'No summary available'}), 404


@app.route('/document/<int:doc_id>/smart-tags')
@login_required
def get_smart_tags(doc_id):
    """Get AI-generated smart tags for a document."""
    document = Document.query.filter_by(id=doc_id, user_id=current_user.id).first_or_404()
    
    if document.ai_tags:
        tags = [tag.strip() for tag in document.ai_tags.split(',')]
        return jsonify({'success': True, 'tags': tags})
    
    # Generate tags if not exists
    if document.extracted_text:
        smart_tags = generate_smart_tags(document.extracted_text, document.subject)
        if smart_tags:
            document.ai_tags = ', '.join(smart_tags)
            db.session.commit()
            return jsonify({'success': True, 'tags': smart_tags})
    
    return jsonify({'success': False, 'error': 'No tags available'}), 404


@app.route('/document/<int:doc_id>/extracted-text')
@login_required
def get_extracted_text(doc_id):
    """Get extracted text from document."""
    document = Document.query.filter_by(id=doc_id, user_id=current_user.id).first_or_404()
    
    if document.extracted_text:
        return jsonify({
            'success': True,
            'text': document.extracted_text,
            'length': len(document.extracted_text)
        })
    
    return jsonify({'success': False, 'error': 'No extracted text available'}), 404


@app.route('/search')
@login_required
def search_documents():
    """Full-text search across documents."""
    query = request.args.get('q', '').strip()
    
    if not query:
        return render_template('search.html', query='', results=[])
    
    # Perform full-text search
    results = search_documents_fulltext(query, current_user.id)
    
    return render_template('search.html', query=query, results=results)


@app.route('/api/search')
@login_required
def api_search_documents():
    """API endpoint for document search."""
    query = request.args.get('q', '').strip()
    
    if not query:
        return jsonify({'success': False, 'error': 'Query required'}), 400
    
    results = search_documents_fulltext(query, current_user.id)
    
    results_list = []
    for doc in results:
        results_list.append({
            'id': doc.id,
            'original_filename': doc.original_filename,
            'year': doc.year,
            'subject': doc.subject,
            'summary': doc.summary,
            'upload_date': doc.upload_date.strftime('%Y-%m-%d')
        })
    
    return jsonify({'success': True, 'results': results_list, 'count': len(results_list)})


@app.route('/document/<int:doc_id>/recommendations')
@login_required
def document_recommendations(doc_id):
    """Get recommended documents based on similarity."""
    document = Document.query.filter_by(id=doc_id, user_id=current_user.id).first_or_404()
    
    recommendations = get_document_recommendations(doc_id)
    
    recommendations_list = []
    for doc in recommendations:
        recommendations_list.append({
            'id': doc.id,
            'original_filename': doc.original_filename,
            'year': doc.year,
            'subject': doc.subject,
            'summary': doc.summary,
            'thumbnail_url': url_for('thumbnail_file', filename=doc.thumbnail_filename) if doc.thumbnail_filename else None
        })
    
    return jsonify({'success': True, 'recommendations': recommendations_list})


@app.route('/ai-features')
@login_required
def ai_features_page():
    """AI features showcase page."""
    # Get some stats
    total_docs = Document.query.filter_by(user_id=current_user.id).count()
    analyzed_docs = Document.query.filter_by(user_id=current_user.id).filter(Document.extracted_text.isnot(None)).count()
    summarized_docs = Document.query.filter_by(user_id=current_user.id).filter(Document.summary.isnot(None)).count()
    
    # Get the actual analyzed documents with summaries
    analyzed_documents = Document.query.filter_by(user_id=current_user.id).filter(
        Document.summary.isnot(None)
    ).order_by(Document.last_analyzed.desc()).limit(6).all()
    
    stats = {
        'total_documents': total_docs,
        'analyzed_documents': analyzed_docs,
        'summarized_documents': summarized_docs,
        'ai_enabled': gemini_model is not None or openai_client is not None,
        'ocr_enabled': os.path.exists(app.config['TESSERACT_CMD'])
    }
    
    return render_template('ai_features.html', stats=stats, analyzed_docs=analyzed_documents)


# ============================================================================
# AI STUDY ASSISTANT CHATBOT ROUTES
# ============================================================================

@app.route('/chat')
@login_required
def chat_index():
    """Main chat interface - shows all chat sessions."""
    sessions = ChatSession.query.filter_by(user_id=current_user.id).order_by(ChatSession.updated_at.desc()).all()
    return render_template('chat/index.html', sessions=sessions)


@app.route('/chat/new', methods=['POST'])
@login_required
def chat_new():
    """Create a new chat session."""
    data = request.get_json()
    document_id = data.get('document_id')
    title = data.get('title', 'New Chat')
    
    # Validate document if provided
    document = None
    if document_id:
        document = Document.query.filter_by(id=document_id, user_id=current_user.id).first()
        if not document:
            return jsonify({'success': False, 'error': 'Document not found'}), 404
        title = f"Chat about {document.original_filename}"
    
    # Create new session
    session = ChatSession(
        user_id=current_user.id,
        document_id=document_id,
        title=title
    )
    db.session.add(session)
    db.session.commit()
    
    return jsonify({
        'success': True,
        'session_id': session.id,
        'title': session.title,
        'redirect_url': url_for('chat_session', session_id=session.id)
    })


@app.route('/chat/<int:session_id>')
@login_required
def chat_session(session_id):
    """View a specific chat session."""
    session = ChatSession.query.filter_by(id=session_id, user_id=current_user.id).first_or_404()
    messages = session.messages.all()
    
    # Get document context if available
    document = session.document if session.document_id else None
    
    return render_template('chat/session.html', 
                         session=session, 
                         messages=messages,
                         document=document)


@app.route('/chat/<int:session_id>/message', methods=['POST'])
@login_required
def chat_send_message(session_id):
    """Send a message in a chat session and get AI response."""
    session = ChatSession.query.filter_by(id=session_id, user_id=current_user.id).first_or_404()
    
    data = request.get_json()
    user_message = data.get('message', '').strip()
    
    if not user_message:
        return jsonify({'success': False, 'error': 'Message cannot be empty'}), 400
    
    # Save user message
    user_msg = ChatMessage(
        session_id=session.id,
        role='user',
        content=user_message
    )
    db.session.add(user_msg)
    
    # Generate AI response
    try:
        # Build context from document if available
        context = ""
        if session.document:
            doc = session.document
            context = f"""Document Context:
Filename: {doc.original_filename}
Subject: {doc.subject}
Year: {doc.year}
"""
            if doc.summary:
                context += f"Summary: {doc.summary}\n"
            if doc.extracted_text:
                # Use first 3000 characters of extracted text
                context += f"\nDocument Content (excerpt):\n{doc.extracted_text[:3000]}...\n"
        
        # Get conversation history
        previous_messages = session.messages.order_by(ChatMessage.created_at).all()
        conversation_history = []
        for msg in previous_messages[-10:]:  # Last 10 messages for context
            conversation_history.append({
                'role': msg.role,
                'parts': [msg.content]
            })
        
        # Add current user message
        conversation_history.append({
            'role': 'user',
            'parts': [user_message]
        })
        
        # Call Gemini API
        if gemini_model:
            system_prompt = f"""You are an AI Study Assistant helping students understand their documents and study materials.
You are helpful, friendly, and focused on education.

{context}

Answer the student's questions based on the document context provided. If asked to generate quizzes, create multiple-choice questions. If asked for study plans, provide structured schedules."""
            
            chat = gemini_model.start_chat(history=conversation_history[:-1])
            response = chat.send_message(user_message)
            ai_response = response.text
        else:
            ai_response = "AI is not configured. Please set up Gemini API key in .env file."
        
        # Save AI response
        ai_msg = ChatMessage(
            session_id=session.id,
            role='assistant',
            content=ai_response
        )
        db.session.add(ai_msg)
        
        # Update session timestamp
        session.updated_at = datetime.utcnow()
        
        db.session.commit()
        
        return jsonify({
            'success': True,
            'user_message': {
                'id': user_msg.id,
                'content': user_msg.content,
                'created_at': user_msg.created_at.isoformat()
            },
            'ai_message': {
                'id': ai_msg.id,
                'content': ai_msg.content,
                'created_at': ai_msg.created_at.isoformat()
            }
        })
        
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/chat/<int:session_id>/delete', methods=['POST'])
@login_required
def chat_delete(session_id):
    """Delete a chat session."""
    session = ChatSession.query.filter_by(id=session_id, user_id=current_user.id).first_or_404()
    db.session.delete(session)
    db.session.commit()
    flash('Chat session deleted successfully', 'success')
    return redirect(url_for('chat_index'))


@app.route('/chat/quiz/generate', methods=['POST'])
@login_required
def generate_quiz():
    """Generate quiz questions from a document using AI."""
    data = request.get_json()
    document_id = data.get('document_id')
    num_questions = data.get('num_questions', 5)
    difficulty = data.get('difficulty', 'medium')  # easy, medium, hard
    
    if not document_id:
        return jsonify({'success': False, 'error': 'Document ID required'}), 400
    
    document = Document.query.filter_by(id=document_id, user_id=current_user.id).first()
    if not document:
        return jsonify({'success': False, 'error': 'Document not found'}), 404
    
    if not document.extracted_text:
        return jsonify({'success': False, 'error': 'Document has no extracted text. Please analyze it first.'}), 400
    
    try:
        if gemini_model:
            prompt = f"""Based on the following document content, generate {num_questions} multiple-choice questions at {difficulty} difficulty level.

Document: {document.original_filename}
Subject: {document.subject}

Content:
{document.extracted_text[:4000]}

Generate questions in this exact JSON format:
{{
  "questions": [
    {{
      "question": "Question text here?",
      "options": ["Option A", "Option B", "Option C", "Option D"],
      "correct_answer": 0,
      "explanation": "Why this answer is correct"
    }}
  ]
}}

Make the questions educational and test real understanding, not just memorization."""
            
            response = gemini_model.generate_content(prompt)
            quiz_text = response.text
            
            # Try to extract JSON from response
            import json
            import re
            
            # Find JSON in response
            json_match = re.search(r'\{[\s\S]*\}', quiz_text)
            if json_match:
                quiz_data = json.loads(json_match.group())
            else:
                # If no JSON found, return raw text
                return jsonify({
                    'success': True,
                    'raw_response': quiz_text
                })
            
            return jsonify({
                'success': True,
                'quiz': quiz_data,
                'document': {
                    'id': document.id,
                    'filename': document.original_filename
                }
            })
        else:
            return jsonify({'success': False, 'error': 'AI not configured'}), 500
            
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/chat/study-plan/generate', methods=['POST'])
@login_required
def generate_study_plan():
    """Generate a personalized study plan based on user's documents."""
    data = request.get_json()
    goal = data.get('goal', 'General study improvement')
    duration_days = data.get('duration_days', 7)
    hours_per_day = data.get('hours_per_day', 2)
    
    # Get user's documents
    documents = Document.query.filter_by(user_id=current_user.id).order_by(Document.upload_date.desc()).limit(20).all()
    
    if not documents:
        return jsonify({'success': False, 'error': 'No documents found. Upload some study materials first.'}), 400
    
    try:
        if gemini_model:
            # Build document list
            doc_list = ""
            for doc in documents:
                doc_list += f"- {doc.original_filename} (Subject: {doc.subject}, Year: {doc.year})\n"
                if doc.summary:
                    doc_list += f"  Summary: {doc.summary[:200]}...\n"
            
            prompt = f"""Create a personalized {duration_days}-day study plan for a student with the following goal: {goal}

Study Schedule:
- Duration: {duration_days} days
- Study time per day: {hours_per_day} hours

Available Study Materials:
{doc_list}

Generate a structured study plan in JSON format:
{{
  "title": "Your Study Plan Title",
  "overview": "Brief overview paragraph",
  "daily_schedule": [
    {{
      "day": 1,
      "focus": "Topic/Subject",
      "tasks": [
        "Task 1 description",
        "Task 2 description"
      ],
      "documents": ["filename1.pdf", "filename2.pdf"],
      "estimated_hours": 2
    }}
  ],
  "tips": [
    "Study tip 1",
    "Study tip 2"
  ]
}}

Make it realistic, achievable, and motivating!"""
            
            response = gemini_model.generate_content(prompt)
            plan_text = response.text
            
            # Try to extract JSON
            import json
            import re
            
            json_match = re.search(r'\{[\s\S]*\}', plan_text)
            if json_match:
                plan_data = json.loads(json_match.group())
            else:
                return jsonify({
                    'success': True,
                    'raw_response': plan_text
                })
            
            return jsonify({
                'success': True,
                'study_plan': plan_data
            })
        else:
            return jsonify({'success': False, 'error': 'AI not configured'}), 500
            
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


if __name__ == '__main__':
    with app.app_context():
        db.create_all()
    app.run(debug=True)
